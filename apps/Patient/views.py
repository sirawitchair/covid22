import os
import xlsxwriter
from datetime import date, timedelta
from io import StringIO, BytesIO

from datetime import date
from django.urls import reverse_lazy
from django.views.generic.edit import CreateView
from django.views.generic import ListView
from django.contrib.auth.mixins import LoginRequiredMixin
from django.core.exceptions import PermissionDenied
from django.http import HttpResponse,HttpResponseForbidden
from django.conf import settings
from django.shortcuts import get_object_or_404,render, redirect

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH
from openpyxl import Workbook,load_workbook
from openpyxl.styles import Font


from .models import Patient
from .forms import PatientForm

from apps.modules.utils import thai_date


class PatientCreateView(LoginRequiredMixin, CreateView):
    model = Patient
    form_class = PatientForm
    template_name = 'Patient/patient_create_view.html'
    login_url = '/login/'
    success_url = reverse_lazy('Patient:List')

    def get_initial(self):
        return {'unit': self.request.user.unit }

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        context["title"] = "เพิ่มข้อมูล"
        return context
    
    def form_valid(self, form):
        instance = form.save(commit=False)
        instance.date = date.today()
        instance.data_user = self.request.user      
        instance.save()  
        return super().form_valid(form)


class PatientListView(LoginRequiredMixin, ListView):
    model = Patient
    template_name = 'Patient/patient_list_view.html'
    login_url = '/login/'
    paginate_by = 30

    def get_queryset(self) :
        currect_user = self.request.user
        user_groups = [g.name for g in currect_user.groups.all()]
        qs = Patient.objects.all()
        # if 'DCAUser' in user_groups or 'DOUser' in  user_groups:            
        #     qs = Patient.objects.all()
        # elif 'MedicalUser' in user_groups:
        #     qs = Patient.objects.filter(document_approved = True)
        # elif 'UnitAdmin' in user_groups:            
        #     qs = Patient.objects.filter(unit = currect_user.unit)
        # else:            
        #     qs = Patient.objects.filter(data_user = currect_user)

        qs = qs.order_by('-date')
        return qs
    
def RadioWorkbook(request):
    xls_unit= Patient.objects.all()
    font=Font(name='TH SarabunPSK',size=18)
    testxls =  os.path.join(settings.BASE_DIR,'apps/Patient/templates/Patient/documents/radioexcel.xlsx')
    workbook = load_workbook(filename=testxls)
    xls_title= f"Data.xlsx"
    sheet_number = 0
    sheets = workbook.sheetnames
    sheet = workbook[sheets[sheet_number]]
    
    for row_num , person in enumerate(xls_unit,2):
        print(person)
        sheet[f"A{row_num}"] = row_num-1
        sheet[f"B{row_num}"] = str(person.date)
        sheet[f"C{row_num}"] = person.full_name
        sheet[f"D{row_num}"] = person.mobile
        sheet[f"E{row_num}"] = str(person.unit)
        sheet[f"F{row_num}"] = str(person.data_user)
        sheet[f"G{row_num}"] = 'อนุมัติแล้ว' if person.document_approved else 'ยังไม่ได้รับการอนุมัติ'
        
    f = BytesIO()
    workbook.save(f)
    length = f.tell()
    f.seek(0)
    response = HttpResponse(
        f.getvalue(),
        content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
    )
    # print('xls = ',xls_title)
    response['Content-Disposition'] = 'attachment; filename="{}"'.format(xls_title)
    response['Content-Length'] = length
    return response

def RadioDocument(request, patient_id):
    patient = get_object_or_404(Patient, pk=patient_id)
    radio_document =  os.path.join(settings.BASE_DIR,'apps/Patient/templates/Patient/documents/radio.docx')
    print(radio_document)
    docx_title= f"radio_report.docx"
    document = Document(radio_document)
    
    

    if patient.atk_date:
        atk_date = "ATK เมื่อ " + thai_date(patient.atk_date)
    else:
        atk_date = ""
    
    if patient.rt_pcr:
        if patient.rt_pcr_place:
            rt_pcr_data = "ตรวจ RT PCR เมื่อ " + thai_date(patient.rt_pcr) + " ที่ " + patient.rt_pcr_place
        else:
            rt_pcr_data = "ตรวจ RT PCR เมื่อ " + thai_date(patient.rt_pcr)
    else:
        rt_pcr_data = ""
    
    dic = {            
            'full_name': f"{patient.get_rank_display()} {patient.full_name}" ,
            'unit': patient.unit.short_name,
            'address' : patient.address,
            'emergency_name' : patient.emergency_name,
            'emergency_mobile' : patient.emergency_mobile,
            'mobile' : patient.mobile,
            'symptom' : patient.symptom,
            'atk_date' : atk_date,
            'rt_pcr_data': rt_pcr_data,
            'treatment' : patient.get_treatment_display()
        }

    for para in document.paragraphs:
        # print('para.text = ',para.text)
        for key, value in dic.items():
            # print(key,' in para.text = ',key in para.text)
            if key in para.text:
                inline = para.runs
                for i in range(len(inline)):
                    if key in inline[i].text:
                        if value: # ถ้ามีการกรอกข้อมูล
                            Value = value
                        else: # ถ้าไม่มีการกรอกข้อมูล
                            Value = ""
                        text = inline[i].text.replace(key, Value)
                        inline[i].text = text
                # new_para = insert_paragraph_after(para,"test insert paragraph\ttab character\ta\tb", 'word_from04_header')
                # tab_stops = new_para.paragraph_format.tab_stops
                # tab_stop = tab_stops.add_tab_stop(Cm(8), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.DOTS)
                # tab_stop = tab_stops.add_tab_stop(Cm(10))

    table = document.tables[0]
    # row0 = t.rows[0] # for example
    # row1 = t.rows[-1]
    # row0._tr.addnext(row1._tr)
    table.cell(1, 0).text = "จาก (From)  " + patient.unit.short_name
    # paragraph = table.cell(0, 2).add_paragraph("หมู่ วัน - เวลา\n" + thai_date(date.today()))
    # paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    table.cell(0, 2).text = "หมู่ วัน - เวลา\n" + thai_date(date.today())
    table.cell(0, 2).alignment = WD_ALIGN_PARAGRAPH.CENTER
    # table.cell(1, 1).text = dic["FullName"]
    table.rows[1].style = document.styles['Normal']

    # coresident_table = document.tables[1]
    # for (i, cs) in enumerate(home_owner.CoResident.all()):        
    #     coresident_table.cell(2 + i, 0).text = cs.full_name
    #     coresident_table.cell(2 + i, 1).text = str(cs.birth_day)
    #     coresident_table.cell(2 + i, 2).text = cs.get_relation_display()
    #     # coresident_table.cell(2 + i, 3).style = document.styles['NormalText']
    #     coresident_table.rows[2 + i].style = document.styles['NormalText']
    
    # vehical_table = document.tables[2]
    # for (i, vh) in enumerate(home_owner.HomeParker.all()):        
    #     vehical_table.cell(1 + i, 0).text = vh.get_type_display()
    #     vehical_table.cell(1 + i, 1).text = vh.brand
    #     vehical_table.cell(1 + i, 2).text = vh.color
    #     vehical_table.cell(1 + i, 3).text = vh.plate
    #     vehical_table.cell(1 + i, 4).text = vh.province
    #     # vehical_table.cell(1 + i, 3).style = document.styles['NormalText']
    #     vehical_table.rows[1 + i].style = document.styles['NormalText']

    f = BytesIO()
    document.save(f)
    length = f.tell()
    f.seek(0)
    response = HttpResponse(
        f.getvalue(),
        content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )
    # print('docx_title = ',docx_title)
    response['Content-Disposition'] = 'attachment; filename="{}"'.format(docx_title)
    response['Content-Length'] = length
    return response

def updateDocument(request,patient_id):
    
    patient_update = Patient.objects.get(id=patient_id)
    if patient_update.document_approved:
        return HttpResponseForbidden()
    if patient_update.data_user!=request.user:
        return HttpResponseForbidden()
    form=PatientForm(instance=patient_update)
    if request.method == 'POST':
        form=PatientForm(request.POST,instance=patient_update)
        if form.is_valid():
            form.save()
            return redirect('/')
    context = {"form" : form , "title" : "แก้ไขข้อมูล"}
    return render(request, 'Patient/patient_create_view.html', context)
